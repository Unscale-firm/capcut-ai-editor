"""
Export the FINAL cut's script (what is actually spoken in the finished video) as .docx and .pdf.

Timestamps are positions in the finished, sped-up video, so you can jump straight to a line.
[SIDE] marks the passages that play on the side camera.

Run:
  venv/Scripts/python.exe pipeline/script_doc.py --work work_vsl --out C:/Users/User/Downloads/vsl-script
"""
import os, json, argparse, sys
sys.path.insert(0, os.path.dirname(__file__))
from build_vsl import SIDE

TITLE = "UNSCALE VSL — Final Script"
ORANGE = (0xE8, 0x76, 0x2D)


def load(work, speed):
    words = json.load(open(os.path.join(work, "words_cut.json"), encoding="utf-8"))
    sents, cur = [], []
    for x in words:
        cur.append(x)
        if x["word"].rstrip().endswith((".", "?", "!")):
            sents.append(cur)
            cur = []
    if cur:
        sents.append(cur)

    on_side = lambda t: any(a <= t < b for a, b in SIDE)
    blocks, para, start, side = [], [], None, None
    for s in sents:
        t, sd = s[0]["start"], on_side(s[0]["start"])
        txt = " ".join(x["word"] for x in s)
        if start is None:
            start, side = t, sd
        elif sd != side or len(" ".join(para)) > 420:
            blocks.append((start / speed, side, " ".join(para)))
            para, start, side = [], t, sd
        para.append(txt)
    if para:
        blocks.append((start / speed, side, " ".join(para)))
    return blocks


def stamp(t):
    return f"{int(t // 60):02d}:{int(t % 60):02d}"


def to_docx(blocks, path, meta):
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name, st.font.size = "Calibri", Pt(11)

    doc.add_heading(TITLE, 0)
    for line in meta:
        doc.add_paragraph(line)
    doc.add_paragraph("Timestamps are positions in the FINAL video. [SIDE] = cuts to the side camera.")
    doc.add_paragraph()

    for t, side, text in blocks:
        h = doc.add_paragraph()
        r = h.add_run(stamp(t))
        r.bold = True
        r.font.color.rgb = RGBColor(*ORANGE)
        r.font.size = Pt(12)
        if side:
            s = h.add_run("   [SIDE]")
            s.bold = True
            s.font.size = Pt(9)
            s.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(14)
    doc.save(path)
    print("wrote", path)


def to_pdf(blocks, path, meta):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    ss = getSampleStyleSheet()
    h_st = ParagraphStyle("h", parent=ss["Heading1"], fontSize=18, spaceAfter=10)
    m_st = ParagraphStyle("m", parent=ss["Normal"], fontSize=9, textColor=HexColor("#666666"),
                          spaceAfter=2)
    t_st = ParagraphStyle("t", parent=ss["Normal"], fontSize=11, textColor=HexColor("#E8762D"),
                          spaceBefore=10, spaceAfter=3, fontName="Helvetica-Bold")
    b_st = ParagraphStyle("b", parent=ss["Normal"], fontSize=10.5, leading=15, spaceAfter=6)

    doc = SimpleDocTemplate(path, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                            leftMargin=2 * cm, rightMargin=2 * cm, title=TITLE)
    flow = [Paragraph(TITLE, h_st)]
    for line in meta:
        flow.append(Paragraph(line, m_st))
    flow.append(Paragraph("Timestamps are positions in the FINAL video. [SIDE] = cuts to the side camera.", m_st))
    flow.append(Spacer(1, 10))

    for t, side, text in blocks:
        head = stamp(t) + ('&nbsp;&nbsp;&nbsp;<font size="8" color="#888888">[SIDE]</font>' if side else "")
        flow.append(Paragraph(head, t_st))
        flow.append(Paragraph(text.replace("&", "&amp;").replace("<", "&lt;"), b_st))
    doc.build(flow)
    print("wrote", path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--work", default="work_vsl")
    ap.add_argument("--speed", type=float, default=1.15)
    ap.add_argument("--out", required=True, help="path WITHOUT extension")
    a = ap.parse_args()

    blocks = load(a.work, a.speed)
    meta = ["Final video: vsl-captioned.mp4 — 18:20 — 1920x1080 — 1.15x speed",
            "Source: 30:33 raw → 21:06 cut → 18:20 sped up",
            f"{len(blocks)} blocks · 22 angle-switch windows"]
    to_docx(blocks, a.out + ".docx", meta)
    to_pdf(blocks, a.out + ".pdf", meta)


if __name__ == "__main__":
    main()
